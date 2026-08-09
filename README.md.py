import streamlit as st
import pandas as pd
import sqlite3
from datetime import datetime, date, timedelta
from pathlib import Path
from io import BytesIO
import hashlib
import secrets

try:
    import cv2
    import numpy as np
except Exception:
    cv2 = None
    np = None

try:
    import zxingcpp
except Exception:
    zxingcpp = None

try:
    from docx import Document
except Exception:
    Document = None

# =========================================================
# CONFIGURATION
# =========================================================

APP_DIR = Path(__file__).parent
DB_PATH = APP_DIR / "sales_app.db"
UPLOAD_DIR = APP_DIR / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

st.set_page_config(
    page_title="SalesFlow | إدارة المبيعات والمخزون",
    page_icon="📦",
    layout="wide",
    initial_sidebar_state="expanded",
)

# =========================================================
# CUSTOM STYLE
# =========================================================

st.markdown(
    """
    <style>

    .main {
        background: #f8fafc;
    }

    .block-container {
        padding-top: 1.5rem;
        max-width: 1450px;
    }

    [data-testid="stSidebar"] {
        background: #0f172a;
    }

    [data-testid="stSidebar"] * {
        color: #e2e8f0 !important;
    }

    .app-title {
        font-size: 2.1rem;
        font-weight: 800;
        color: #0f172a;
        margin-bottom: 5px;
    }

    .app-subtitle {
        color: #64748b;
        margin-bottom: 25px;
    }

    .kpi {
        background: white;
        border: 1px solid #e2e8f0;
        border-radius: 16px;
        padding: 20px;
        box-shadow: 0 4px 15px rgba(15, 23, 42, 0.05);
        min-height: 125px;
    }

    .kpi-label {
        color: #64748b;
        font-size: 0.9rem;
    }

    .kpi-value {
        color: #0f172a;
        font-size: 1.8rem;
        font-weight: 800;
        margin-top: 7px;
    }

    .small {
        color: #64748b;
        font-size: 0.82rem;
    }

    .section {
        background: white;
        border: 1px solid #e2e8f0;
        border-radius: 16px;
        padding: 20px;
        margin-bottom: 18px;
    }

    .sf-table-wrap {
        width: 100%;
        overflow-x: auto;
        border: 1px solid #e2e8f0;
        border-radius: 14px;
        background: white;
        margin: 10px 0 18px 0;
    }

    .sf-table {
        width: 100%;
        border-collapse: collapse;
        font-size: 0.92rem;
        direction: rtl;
    }

    .sf-table th {
        background: #f1f5f9;
        color: #0f172a;
        font-weight: 800;
        padding: 12px 14px;
        border-bottom: 1px solid #e2e8f0;
        white-space: nowrap;
        text-align: right;
    }

    .sf-table td {
        padding: 11px 14px;
        border-bottom: 1px solid #eef2f7;
        color: #334155;
        text-align: right;
        vertical-align: middle;
    }

    .sf-table tr:last-child td {
        border-bottom: 0;
    }

    .sf-table tr:hover td {
        background: #f8fafc;
    }

    .action-card {
        background: linear-gradient(135deg, #0f172a, #1e293b);
        color: white;
        border-radius: 18px;
        padding: 20px;
        margin: 12px 0;
    }

    </style>
    """,
    unsafe_allow_html=True,
)

# =========================================================
# DATABASE
# =========================================================


def get_db():
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.execute("PRAGMA foreign_keys=ON")
    return conn


def init_db():

    conn = get_db()
    cursor = conn.cursor()

    cursor.executescript(
        """

        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL UNIQUE,
            role TEXT NOT NULL CHECK(role IN ('manager', 'deputy', 'employee')),
            pin_hash TEXT NOT NULL,
            rep_id INTEGER,
            active INTEGER DEFAULT 1,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(rep_id) REFERENCES reps(id)
        );

        CREATE TABLE IF NOT EXISTS activity_log (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            user_name TEXT NOT NULL,
            role TEXT NOT NULL,
            action TEXT NOT NULL,
            details TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        );

        CREATE TABLE IF NOT EXISTS reps (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            phone TEXT,
            active INTEGER DEFAULT 1
        );

        CREATE TABLE IF NOT EXISTS customers (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            area TEXT,
            rep_id INTEGER,
            active INTEGER DEFAULT 1,
            FOREIGN KEY(rep_id) REFERENCES reps(id)
        );

        CREATE TABLE IF NOT EXISTS products (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            sku TEXT UNIQUE NOT NULL,
            name TEXT NOT NULL,
            category TEXT,
            min_stock INTEGER DEFAULT 0,
            dormant_days INTEGER DEFAULT 30,
            active INTEGER DEFAULT 1
        );

        CREATE TABLE IF NOT EXISTS rep_stock (
            rep_id INTEGER,
            product_id INTEGER,
            qty REAL DEFAULT 0,

            PRIMARY KEY(rep_id, product_id),

            FOREIGN KEY(rep_id)
                REFERENCES reps(id),

            FOREIGN KEY(product_id)
                REFERENCES products(id)
        );

        CREATE TABLE IF NOT EXISTS visits (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            rep_id INTEGER NOT NULL,
            customer_id INTEGER NOT NULL,
            visit_date TEXT NOT NULL,
            notes TEXT,
            image_path TEXT,

            FOREIGN KEY(rep_id)
                REFERENCES reps(id),

            FOREIGN KEY(customer_id)
                REFERENCES customers(id)
        );

        CREATE TABLE IF NOT EXISTS visit_stock (
            visit_id INTEGER,
            product_id INTEGER,
            qty REAL DEFAULT 0,

            PRIMARY KEY(visit_id, product_id),

            FOREIGN KEY(visit_id)
                REFERENCES visits(id),

            FOREIGN KEY(product_id)
                REFERENCES products(id)
        );

        CREATE TABLE IF NOT EXISTS orders (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            rep_id INTEGER NOT NULL,
            customer_id INTEGER NOT NULL,
            order_date TEXT NOT NULL,
            invoice_no TEXT,
            image_path TEXT,

            FOREIGN KEY(rep_id)
                REFERENCES reps(id),

            FOREIGN KEY(customer_id)
                REFERENCES customers(id)
        );

        CREATE TABLE IF NOT EXISTS order_items (
            order_id INTEGER,
            product_id INTEGER,
            qty REAL DEFAULT 0,

            PRIMARY KEY(order_id, product_id),

            FOREIGN KEY(order_id)
                REFERENCES orders(id),

            FOREIGN KEY(product_id)
                REFERENCES products(id)
        );

        CREATE TABLE IF NOT EXISTS targets (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            rep_id INTEGER,
            month TEXT NOT NULL,
            target_sales REAL DEFAULT 0,
            target_collection REAL DEFAULT 0
        );

        CREATE TABLE IF NOT EXISTS collections (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            rep_id INTEGER,
            customer_id INTEGER,
            collection_date TEXT,
            amount REAL DEFAULT 0,
            reference TEXT
        );

        """
    )

    # -----------------------------------------------------
    # Company workflow migrations
    # -----------------------------------------------------
    def add_column_if_missing(table, column, definition):
        cols = [r[1] for r in cursor.execute(f"PRAGMA table_info({table})").fetchall()]
        if column not in cols:
            cursor.execute(f"ALTER TABLE {table} ADD COLUMN {column} {definition}")

    add_column_if_missing("customers", "phone", "TEXT")
    add_column_if_missing("customers", "notes", "TEXT")
    add_column_if_missing("products", "unit_price", "REAL DEFAULT 0")
    add_column_if_missing("products", "barcode", "TEXT")
    add_column_if_missing("orders", "invoice_barcode", "TEXT")
    add_column_if_missing("orders", "barcode_image_path", "TEXT")
    add_column_if_missing("orders", "total_amount", "REAL DEFAULT 0")
    add_column_if_missing("order_items", "unit_price", "REAL DEFAULT 0")

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS weekly_reports (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            rep_id INTEGER NOT NULL,
            week_start TEXT NOT NULL,
            week_end TEXT NOT NULL,
            title TEXT NOT NULL,
            summary TEXT,
            file_path TEXT NOT NULL,
            status TEXT DEFAULT 'pending' CHECK(status IN ('pending','approved','rejected')),
            manager_note TEXT,
            reviewed_by INTEGER,
            reviewed_at TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(rep_id, week_start, week_end),
            FOREIGN KEY(rep_id) REFERENCES reps(id)
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS stock_movements (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            rep_id INTEGER NOT NULL,
            product_id INTEGER NOT NULL,
            movement_type TEXT NOT NULL,
            qty REAL NOT NULL,
            reference TEXT,
            image_path TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(rep_id) REFERENCES reps(id),
            FOREIGN KEY(product_id) REFERENCES products(id)
        )
    """)

    # Convert the previous generic demo catalog to paper-company products.
    old_demo = cursor.execute("SELECT COUNT(*) FROM products WHERE sku IN ('SKU-001','SKU-002','SKU-003','SKU-004')").fetchone()[0]
    if old_demo == 4:
        cursor.execute("DELETE FROM products")
        cursor.executemany("""
            INSERT INTO products(sku,name,category,min_stock,dormant_days,unit_price,barcode)
            VALUES (?,?,?,?,?,?,?)
        """, [
            ("PAPER-A4-80", "ورق A4 80 جرام - كرتون", "ورق", 5, 30, 85.0, "6281000000011"),
            ("PAPER-A3-80", "ورق A3 80 جرام - كرتون", "ورق", 3, 30, 125.0, "6281000000012"),
            ("PAPER-A4-70", "ورق A4 70 جرام - كرتون", "ورق", 5, 30, 78.0, "6281000000013"),
            ("PAPER-COLOR", "ورق ملون - كرتون", "ورق", 2, 45, 110.0, "6281000000014"),
        ])

    # -----------------------------------------------------
    # Demo data
    # -----------------------------------------------------

    if cursor.execute("SELECT COUNT(*) FROM reps").fetchone()[0] == 0:

        cursor.execute(
            """
            INSERT INTO reps(name, phone)
            VALUES (?, ?)
            """,
            ("مندوب تجريبي", "0500000000"),
        )

    if cursor.execute("SELECT COUNT(*) FROM customers").fetchone()[0] == 0:

        rep_id = cursor.execute(
            "SELECT id FROM reps LIMIT 1"
        ).fetchone()[0]

        cursor.executemany(
            """
            INSERT INTO customers(name, area, rep_id)
            VALUES (?, ?, ?)
            """,
            [
                ("عميل الرياض 01", "الرياض", rep_id),
                ("عميل الطائف 01", "الطائف", rep_id),
                ("عميل جدة 01", "جدة", rep_id),
            ],
        )

    if cursor.execute("SELECT COUNT(*) FROM products").fetchone()[0] == 0:

        cursor.executemany(
            """
            INSERT INTO products
            (sku, name, category, min_stock, dormant_days, unit_price, barcode)
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            [
                ("PAPER-A4-80", "ورق A4 80 جرام - كرتون", "ورق", 5, 30, 85.0, "6281000000011"),
                ("PAPER-A3-80", "ورق A3 80 جرام - كرتون", "ورق", 3, 30, 125.0, "6281000000012"),
                ("PAPER-A4-70", "ورق A4 70 جرام - كرتون", "ورق", 5, 30, 78.0, "6281000000013"),
                ("PAPER-COLOR", "ورق ملون - كرتون", "ورق", 2, 45, 110.0, "6281000000014"),
            ],
        )

    # -----------------------------------------------------
    # Initial company accounts (PINs are stored as hashes)
    # -----------------------------------------------------
    def pin_hash(pin):
        return hashlib.sha256(pin.encode("utf-8")).hexdigest()

    seed_accounts = [
        ("ناصر علي", "manager", "5321"),
        ("أحمد", "deputy", "5522"),
        ("عمر", "employee", "7418"),
        ("خالد", "employee", "3964"),
        ("فهد", "employee", "8257"),
    ]

    for name, role_name, pin in seed_accounts:
        existing = cursor.execute(
            "SELECT id, rep_id FROM users WHERE name = ?", (name,)
        ).fetchone()
        if existing is None:
            rep_id_for_user = None
            if role_name == "employee":
                rep = cursor.execute("SELECT id FROM reps WHERE name = ?", (name,)).fetchone()
                if rep is None:
                    cursor.execute("INSERT INTO reps(name, phone) VALUES (?, NULL)", (name,))
                    rep_id_for_user = cursor.lastrowid
                else:
                    rep_id_for_user = rep[0]
            cursor.execute(
                "INSERT INTO users(name, role, pin_hash, rep_id) VALUES (?, ?, ?, ?)",
                (name, role_name, pin_hash(pin), rep_id_for_user),
            )

    conn.commit()
    conn.close()


init_db()

# =========================================================
# DATABASE HELPERS
# =========================================================


def query(sql, params=()):

    conn = get_db()

    df = pd.read_sql_query(
        sql,
        conn,
        params=params,
    )

    conn.close()

    return df


def execute(sql, params=()):

    conn = get_db()
    cursor = conn.cursor()
    cursor.execute(sql, params)
    conn.commit()
    last_id = cursor.lastrowid

    # Audit trail: every write operation is associated with the logged-in user.
    if st.session_state.get("authenticated"):
        sql_clean = " ".join(sql.strip().split())
        operation = sql_clean.split(" ")[0].upper() if sql_clean else "WRITE"
        table = "غير محدد"
        upper = sql_clean.upper()
        for candidate in ["USERS", "REPS", "CUSTOMERS", "PRODUCTS", "REP_STOCK", "VISITS", "ORDERS", "TARGETS", "COLLECTIONS", "VISIT_STOCK", "ORDER_ITEMS"]:
            if candidate in upper:
                table = candidate
                break
        user = st.session_state.get("user", {})
        cursor.execute(
            "INSERT INTO activity_log(user_id,user_name,role,action,details) VALUES (?,?,?,?,?)",
            (user.get("id"), user.get("name", "غير معروف"), user.get("role", "unknown"), operation, f"جدول: {table}"),
        )
        conn.commit()

    conn.close()
    return last_id


def hash_pin(pin):
    return hashlib.sha256(str(pin).encode("utf-8")).hexdigest()


def log_activity(action, details=""):
    user = st.session_state.get("user")
    if not user:
        return
    conn = get_db()
    conn.execute(
        "INSERT INTO activity_log(user_id,user_name,role,action,details) VALUES (?,?,?,?,?)",
        (user["id"], user["name"], user["role"], action, details),
    )
    conn.commit()
    conn.close()


def role_label(role):
    return {"manager": "مدير", "deputy": "نائب مدير", "employee": "موظف"}.get(role, role)


def is_admin():
    return st.session_state.get("user", {}).get("role") in {"manager", "deputy"}


def can_manage_accounts():
    return is_admin()


def authenticate():
    if st.session_state.get("authenticated"):
        return True

    st.markdown("## 🔐 تسجيل الدخول")
    st.caption("الدخول يتم بالرمز السري فقط — بدون رقم جوال.")
    with st.form("login_form"):
        users = query("SELECT id, name, role FROM users WHERE active = 1 ORDER BY CASE role WHEN 'manager' THEN 1 WHEN 'deputy' THEN 2 ELSE 3 END, name")
        name = st.selectbox("اختر حسابك", users["name"].tolist()) if len(users) else None
        pin = st.text_input("الرمز السري", type="password", max_chars=20)
        submitted = st.form_submit_button("دخول", use_container_width=True)
        if submitted and name:
            row = query("SELECT id,name,role,rep_id,pin_hash FROM users WHERE name = ? AND active = 1", (name,))
            if len(row) and secrets.compare_digest(hash_pin(pin), row.iloc[0]["pin_hash"]):
                st.session_state.authenticated = True
                st.session_state.user = {
                    "id": int(row.iloc[0]["id"]),
                    "name": row.iloc[0]["name"],
                    "role": row.iloc[0]["role"],
                    "rep_id": int(row.iloc[0]["rep_id"]) if pd.notna(row.iloc[0]["rep_id"]) else None,
                }
                log_activity("تسجيل دخول", "تم الدخول للنظام")
                st.rerun()
            else:
                st.error("الرمز السري غير صحيح.")
    return False


if not authenticate():
    st.stop()


def save_upload(uploaded_file, prefix):

    if uploaded_file is None:
        return None

    timestamp = datetime.now().strftime(
        "%Y%m%d_%H%M%S"
    )

    safe_name = (
        uploaded_file.name
        .replace("/", "_")
        .replace("\\", "_")
    )

    path = (
        UPLOAD_DIR
        / f"{prefix}_{timestamp}_{safe_name}"
    )

    path.write_bytes(
        uploaded_file.getbuffer()
    )

    return str(path)


def decode_barcode(uploaded_file):
    """Decode common barcodes/QRs from a camera image when supported."""
    if uploaded_file is None or cv2 is None or np is None:
        return None, "تعذر القراءة التلقائية؛ اكتب الباركود يدوياً."
    try:
        data = np.frombuffer(uploaded_file.getvalue(), dtype=np.uint8)
        image = cv2.imdecode(data, cv2.IMREAD_COLOR)
        if image is None:
            return None, "الصورة غير صالحة."
        if zxingcpp is not None:
            results = zxingcpp.read_barcodes(image)
            for result in results:
                if result.text:
                    return result.text, str(getattr(result, "format", "BARCODE"))
        detector = cv2.QRCodeDetector()
        value, _, _ = detector.detectAndDecode(image)
        if value:
            return value, "QR"
        return None, "لم يتم العثور على باركود واضح."
    except Exception as exc:
        return None, f"تعذر تحليل الصورة: {exc}"


def make_weekly_report_template(rep_name, week_start, week_end):
    if Document is None:
        return None
    doc = Document()
    doc.add_heading("التقرير الأسبوعي للمندوب", level=1)
    doc.add_paragraph(f"الموظف: {rep_name}")
    doc.add_paragraph(f"الأسبوع: {week_start} إلى {week_end}")
    for heading in [
        "1. ملخص العمل خلال الأسبوع",
        "2. العملاء والزيارات",
        "3. المبيعات والفواتير",
        "4. حركة صناديق الورق والمخزون",
        "5. التحصيل",
        "6. المشاكل والملاحظات",
        "7. خطة الأسبوع القادم",
    ]:
        doc.add_heading(heading, level=2)
        doc.add_paragraph("اكتب التفاصيل هنا...")
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def file_bytes(path):
    try:
        p = Path(str(path))
        return p.read_bytes() if p.exists() else None
    except Exception:
        return None


def role_is_admin(role):
    return role in {"manager", "deputy"}


def manager_only():
    return st.session_state.get("user", {}).get("role") == "manager"


def format_number(value):

    try:
        return f"{float(value):,.0f}"

    except Exception:
        return "0"


def kpi(label, value, note=""):

    st.markdown(
        f"""
        <div class="kpi">

            <div class="kpi-label">
                {label}
            </div>

            <div class="kpi-value">
                {value}
            </div>

            <div class="small">
                {note}
            </div>

        </div>
        """,
        unsafe_allow_html=True,
    )


def show_table(df, **kwargs):
    """Render a static HTML table instead of Streamlit's DataFrame widget.
    This avoids browser-side DataFrame/Arrow dynamic-module failures while
    keeping tables readable and responsive.
    """
    if df is None:
        st.info("لا توجد بيانات.")
        return
    try:
        table_df = df.copy()
    except Exception:
        table_df = pd.DataFrame(df)
    if table_df.empty:
        st.info("لا توجد بيانات لعرضها.")
        return
    html = table_df.to_html(index=False, escape=True, classes="sf-table")
    st.markdown(f'<div class="sf-table-wrap">{html}</div>', unsafe_allow_html=True)


# =========================================================
# SIDEBAR
# =========================================================

st.sidebar.markdown(
    """<h1 style="color:white;">📦 SalesFlow</h1>""",
    unsafe_allow_html=True,
)
st.sidebar.caption("نظام إدارة المبيعات والمخزون")

current_user = st.session_state["user"]
role = current_user["role"]
rep_id = current_user.get("rep_id")

st.sidebar.success(f"{current_user['name']} — {role_label(role)}")

if st.sidebar.button("🚪 تسجيل الخروج", use_container_width=True):
    log_activity("تسجيل خروج", "تم تسجيل الخروج")
    st.session_state.clear()
    st.rerun()


# =========================================================
# NAVIGATION
# =========================================================

if role == "employee":
    pages = [
        "🏠 لوحة المندوب",
        "👥 العملاء والزيارات",
        "🧾 تسجيل أوردر",
        "📦 مخزون المندوب",
        "💰 التحصيل",
        "📊 تقارير المندوب",
    ]
else:
    pages = [
        "🏢 لوحة الإدارة",
        "👥 العملاء والمندوبون",
        "📦 الأصناف والمخزون",
        "🎯 التارجت والتحصيل",
        "⚠️ الأصناف الراكدة",
        "📥 التقارير والتصدير",
        "👁️ سجل نشاط الموظفين",
        "👤 ملف الموظف الكامل",
        "🔐 إدارة الحسابات والصلاحيات",
    ]


page = st.sidebar.radio(
    "التنقل",
    pages,
)


st.sidebar.divider()

st.sidebar.caption(
    "SalesFlow v3.1 • شركة أوراق"
)

# =========================================================
# REP DASHBOARD
# =========================================================

if page == "🏠 لوحة المندوب":

    st.markdown(
        '<div class="app-title">مرحباً بك 👋</div>',
        unsafe_allow_html=True,
    )

    st.markdown(
        '<div class="app-subtitle">'
        "ملخص أدائك ومخزونك اليومي"
        "</div>",
        unsafe_allow_html=True,
    )

    today = date.today().isoformat()

    month = date.today().strftime("%Y-%m")

    sales = query(
        """
        SELECT
            COALESCE(SUM(oi.qty), 0) AS qty

        FROM orders o

        JOIN order_items oi
            ON o.id = oi.order_id

        WHERE o.rep_id = ?

        AND substr(
            o.order_date,
            1,
            7
        ) = ?
        """,
        (rep_id, month),
    )

    visits = query(
        """
        SELECT COUNT(*) AS count
        FROM visits
        WHERE rep_id = ?
        AND visit_date = ?
        """,
        (rep_id, today),
    )

    collections = query(
        """
        SELECT
            COALESCE(SUM(amount), 0) AS amount

        FROM collections

        WHERE rep_id = ?

        AND substr(
            collection_date,
            1,
            7
        ) = ?
        """,
        (rep_id, month),
    )

    target = query(
        """
        SELECT
            COALESCE(
                SUM(target_sales),
                0
            ) AS target

        FROM targets

        WHERE rep_id = ?

        AND month = ?
        """,
        (rep_id, month),
    )

    sales_value = float(
        sales.iloc[0]["qty"]
    )

    sales_money = query("""
        SELECT COALESCE(SUM(oi.qty * COALESCE(oi.unit_price, p.unit_price, 0)),0) AS amount
        FROM orders o
        JOIN order_items oi ON oi.order_id=o.id
        JOIN products p ON p.id=oi.product_id
        WHERE o.rep_id=? AND substr(o.order_date,1,7)=?
    """, (rep_id, month))
    sales_money_value = float(sales_money.iloc[0]["amount"])
    customer_count = query("SELECT COUNT(*) AS count FROM customers WHERE rep_id=? AND active=1", (rep_id,))
    invoice_count = query("SELECT COUNT(*) AS count FROM orders WHERE rep_id=? AND substr(order_date,1,7)=?", (rep_id, month))

    collection_value = float(
        collections.iloc[0]["amount"]
    )

    target_value = float(
        target.iloc[0]["target"]
    )

    achievement = (
        sales_value / target_value * 100
        if target_value
        else 0
    )

    cols = st.columns(5)
    with cols[0]: kpi("عملائي", format_number(customer_count.iloc[0]["count"]), "عملاء نشطون")
    with cols[1]: kpi("زيارات اليوم", format_number(visits.iloc[0]["count"]), "زيارة مسجلة")
    with cols[2]: kpi("المبيعات", f"{sales_money_value:,.0f}", "قيمة مبيعات الشهر")
    with cols[3]: kpi("الفواتير", format_number(invoice_count.iloc[0]["count"]), "هذا الشهر")
    with cols[4]: kpi("التحصيل", f"{collection_value:,.0f}", "خلال الشهر")
    st.info("المبيعات هنا هي قيمة الفواتير حسب أسعار الأصناف المسجلة، وليست راتب الموظف أو عمولته.")

    st.markdown("### 📈 مستوى تحقيق التارجت")

    if target_value:

        st.progress(
            min(
                achievement / 100,
                1,
            )
        )

    else:

        st.info(
            "لم يتم تحديد تارجت لهذا الشهر."
        )

    # -----------------------------------------------------
    # Rep stock
    # -----------------------------------------------------

    inventory = query(
        """
        SELECT

            p.sku,

            p.name,

            COALESCE(
                rs.qty,
                0
            ) AS qty,

            p.min_stock

        FROM products p

        LEFT JOIN rep_stock rs

            ON p.id = rs.product_id

            AND rs.rep_id = ?

        WHERE p.active = 1

        ORDER BY qty ASC
        """,
        (rep_id,),
    )

    st.markdown(
        "### 📦 حالة مخزون المندوب"
    )

    if len(inventory):

        inventory["الحالة"] = inventory.apply(

            lambda row:

            "🔴 نفد"

            if row.qty <= 0

            else (

                "🟠 منخفض"

                if row.qty <= row.min_stock

                else "🟢 متوفر"

            ),

            axis=1,
        )

        show_table(
            inventory[
                [
                    "sku",
                    "name",
                    "qty",
                    "min_stock",
                    "الحالة",
                ]
            ]
        )

    st.markdown("### 🕘 آخر عملياتي")
    my_activity = query("""
        SELECT created_at AS الوقت, action AS العملية, details AS التفاصيل
        FROM activity_log
        WHERE user_id=?
        ORDER BY id DESC
        LIMIT 8
    """, (current_user["id"],))
    show_table(my_activity)


# =========================================================
# CUSTOMER VISITS
# =========================================================

elif page == "👥 العملاء والزيارات":

    st.markdown("## 👥 العملاء والزيارات")

    with st.expander("➕ إضافة عميل جديد", expanded=True):
        with st.form("employee_add_customer"):
            new_customer_name = st.text_input("اسم العميل")
            new_customer_phone = st.text_input("رقم جوال العميل")
            new_customer_area = st.text_input("المدينة / المنطقة")
            new_customer_notes = st.text_area("ملاحظات العميل")
            add_customer = st.form_submit_button("إضافة العميل", use_container_width=True)
            if add_customer:
                if not new_customer_name.strip() or not new_customer_phone.strip():
                    st.error("اسم العميل ورقم الجوال مطلوبان.")
                else:
                    execute("INSERT INTO customers(name,area,rep_id,phone,notes) VALUES (?,?,?,?,?)", (new_customer_name.strip(), new_customer_area.strip(), rep_id, new_customer_phone.strip(), new_customer_notes.strip()))
                    log_activity("إضافة عميل", f"أضاف العميل {new_customer_name.strip()} برقم {new_customer_phone.strip()}")
                    st.success("تمت إضافة العميل وربطه بحسابك.")
                    st.rerun()

    customers = query(
        """
        SELECT *

        FROM customers

        WHERE active = 1

        AND rep_id = ?

        ORDER BY name
        """,
        (rep_id,),
    )

    products = query(
        """
        SELECT *

        FROM products

        WHERE active = 1

        ORDER BY name
        """
    )

    if not len(customers):

        st.warning(
            "لا يوجد عملاء مرتبطون بهذا المندوب."
        )

    else:

        customer_name = st.selectbox(
            "العميل",
            customers["name"].tolist(),
        )

        customer_id = int(
            customers.loc[
                customers["name"]
                == customer_name,
                "id",
            ].iloc[0]
        )

        visit_date = st.date_input(
            "تاريخ الزيارة",
            date.today(),
        )

        notes = st.text_area(
            "ملاحظات الزيارة"
        )

        st.markdown("### 📸 مرفق الزيارة")
        visit_camera = st.camera_input("📷 تصوير المستند/المخزون بالكاميرا")
        image = st.file_uploader(
            "أو ارفع صورة المخزون أو الفاتورة",
            type=["jpg", "jpeg", "png", "webp"],
            key="visit_image_upload",
        )

        st.markdown("### 📦 حصر رصيد الأصناف لدى العميل")


        stock_input = {}

        for _, product in products.iterrows():

            stock_input[int(product.id)] = (
                st.number_input(
                    f"{product.name} — {product.sku}",
                    min_value=0.0,
                    step=1.0,
                    value=0.0,
                    key=f"visit_{product.id}",
                )
            )

        if st.button(
            "💾 حفظ الزيارة وحصر المخزون",
            type="primary",
            use_container_width=True,
        ):

            image_path = save_upload(
                visit_camera or image,
                "visit",
            )

            visit_id = execute(
                """
                INSERT INTO visits
                (
                    rep_id,
                    customer_id,
                    visit_date,
                    notes,
                    image_path
                )

                VALUES (?, ?, ?, ?, ?)
                """,
                (
                    rep_id,
                    customer_id,
                    visit_date.isoformat(),
                    notes,
                    image_path,
                ),
            )

            for product_id, qty in stock_input.items():

                execute(
                    """
                    INSERT INTO visit_stock
                    (
                        visit_id,
                        product_id,
                        qty
                    )

                    VALUES (?, ?, ?)
                    """,
                    (
                        visit_id,
                        product_id,
                        qty,
                    ),
                )

            st.success(
                "تم حفظ الزيارة وحصر المخزون بنجاح."
            )

        # -------------------------------------------------
        # Sell-out
        # -------------------------------------------------

        st.markdown(
            "### 📉 Sell-out بين الزيارات"
        )

        sell_out = query(
            """
            WITH visits_data AS (

                SELECT

                    v.id,

                    v.customer_id,

                    v.visit_date,

                    vs.product_id,

                    vs.qty,

                    LAG(vs.qty)

                    OVER (

                        PARTITION BY
                            v.customer_id,
                            vs.product_id

                        ORDER BY
                            v.visit_date,
                            v.id

                    ) AS previous_qty

                FROM visits v

                JOIN visit_stock vs

                    ON v.id =
                    vs.visit_id

                WHERE v.rep_id = ?

            )

            SELECT

                c.name AS العميل,

                p.name AS الصنف,

                v.visit_date AS تاريخ_الزيارة,

                COALESCE(
                    v.previous_qty,
                    0
                ) AS الرصيد_السابق,

                v.qty AS الرصيد_الحالي,

                CASE

                    WHEN v.previous_qty IS NULL
                    THEN 0

                    WHEN v.previous_qty - v.qty < 0
                    THEN 0

                    ELSE
                        v.previous_qty - v.qty

                END AS sell_out

            FROM visits_data v

            JOIN customers c
                ON c.id = v.customer_id

            JOIN products p
                ON p.id = v.product_id

            ORDER BY
                v.visit_date DESC
            """,
            (rep_id,),
        )

        show_table(
            sell_out,
            use_container_width=True,
            hide_index=True,
        )


# =========================================================
# ORDER ENTRY
# =========================================================

elif page == "🧾 تسجيل أوردر":

    st.markdown(
        "## 🧾 تسجيل Sell-in / أوردر"
    )

    customers = query(
        """
        SELECT *

        FROM customers

        WHERE active = 1

        AND rep_id = ?

        ORDER BY name
        """,
        (rep_id,),
    )

    products = query(
        """
        SELECT *

        FROM products

        WHERE active = 1

        ORDER BY name
        """
    )

    if len(customers) and len(products):

        customer_name = st.selectbox(
            "العميل",
            customers["name"].tolist(),
        )

        customer_id = int(
            customers.loc[
                customers["name"]
                == customer_name,
                "id",
            ].iloc[0]
        )

        invoice_number = st.text_input("رقم الفاتورة / الأوردر")
        barcode_camera = st.camera_input("📷 صوّر باركود الفاتورة")
        barcode_code = st.text_input("باركود الفاتورة (يمكن كتابته يدوياً)")
        if barcode_camera is not None and not barcode_code.strip():
            decoded, fmt = decode_barcode(barcode_camera)
            if decoded:
                barcode_code = decoded
                st.success(f"تمت قراءة الباركود: {decoded} ({fmt})")
            else:
                st.warning("لم تتم قراءة الباركود من الصورة. اكتب الرقم يدوياً إذا لزم.")

        st.markdown("### 📸 صورة الفاتورة")
        invoice_camera = st.camera_input("📷 صوّر الفاتورة مباشرة بالكاميرا")
        invoice_image = st.file_uploader(
            "أو ارفع صورة الفاتورة من الجهاز",
            type=["jpg", "jpeg", "png", "webp"],
            key="invoice_file_upload",
        )

        st.markdown("### 🛒 أصناف الأوردر")
        st.caption("يمكنك تصوير الفاتورة بالكاميرا أو رفع صورتها، ثم تسجيل الأصناف والكمية.")


        order_items = {}

        for _, product in products.iterrows():

            stock = query(
                """
                SELECT
                    COALESCE(qty, 0) AS qty

                FROM rep_stock

                WHERE rep_id = ?

                AND product_id = ?
                """,
                (
                    rep_id,
                    int(product.id),
                ),
            )

            available = (
                float(stock.iloc[0]["qty"])
                if len(stock)
                else 0
            )

            order_items[
                int(product.id)
            ] = st.number_input(
                f"{product.name} — المتاح: {available:g}",
                min_value=0.0,
                max_value=max(
                    available,
                    0,
                ),
                step=1.0,
                key=f"order_{product.id}",
            )

        st.markdown("### 🚀 اعتماد الأوردر")
        st.caption("بعد الضغط على الزر يتم حفظ الفاتورة، تسجيل البيع، وتحديث مخزون الموظف مباشرة.")
        if st.button(
            "🧾 تسجيل الأوردر واعتماده",
            type="primary",
            use_container_width=True,
            key="submit_order",
        ):

            selected_items = {
                pid: qty

                for pid, qty
                in order_items.items()

                if qty > 0
            }

            if not selected_items:

                st.error(
                    "اختر صنفاً واحداً على الأقل."
                )

            else:

                insufficient = []

                for product_id, qty in selected_items.items():

                    stock = query(
                        """
                        SELECT
                            COALESCE(qty, 0) AS qty

                        FROM rep_stock

                        WHERE rep_id = ?

                        AND product_id = ?
                        """,
                        (
                            rep_id,
                            product_id,
                        ),
                    )

                    available = (
                        float(
                            stock.iloc[0]["qty"]
                        )
                        if len(stock)
                        else 0
                    )

                    product_name = products.loc[
                        products["id"]
                        == product_id,
                        "name",
                    ].iloc[0]

                    if qty > available:

                        insufficient.append(
                            f"{product_name}: "
                            f"المطلوب {qty:g} "
                            f"والمتاح {available:g}"
                        )

                if insufficient:

                    st.error(
                        "لا يمكن اعتماد الأوردر:\n\n"
                        + "\n".join(
                            insufficient
                        )
                    )

                else:

                    image_path = save_upload(invoice_camera or invoice_image, "invoice")
                    barcode_image_path = save_upload(barcode_camera, "invoice_barcode")

                    total_amount = 0.0
                    for product_id, qty in selected_items.items():
                        row_price = products.loc[products["id"] == product_id, "unit_price"]
                        unit_price = float(row_price.iloc[0]) if len(row_price) and pd.notna(row_price.iloc[0]) else 0.0
                        total_amount += qty * unit_price

                    order_id = execute(
                        """
                        INSERT INTO orders(rep_id,customer_id,order_date,invoice_no,image_path,invoice_barcode,barcode_image_path,total_amount)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                        """,
                        (rep_id, customer_id, date.today().isoformat(), invoice_number, image_path, barcode_code.strip(), barcode_image_path, total_amount),
                    )

                    for product_id, qty in selected_items.items():

                        execute(
                            """
                            INSERT INTO order_items(order_id,product_id,qty,unit_price)
                            VALUES (?, ?, ?, ?)
                            """,
                            (
                                order_id,
                                product_id,
                                qty,
                                float(products.loc[products["id"] == product_id, "unit_price"].iloc[0]) if pd.notna(products.loc[products["id"] == product_id, "unit_price"].iloc[0]) else 0.0,
                            ),
                        )

                        execute(
                            """
                            INSERT INTO rep_stock
                            (
                                rep_id,
                                product_id,
                                qty
                            )

                            VALUES (?, ?, ?)

                            ON CONFLICT
                            (
                                rep_id,
                                product_id
                            )

                            DO UPDATE SET
                                qty = qty - ?
                            """,
                            (
                                rep_id,
                                product_id,
                                qty,
                                qty,
                            ),
                        )

                    for product_id, qty in selected_items.items():
                        execute("INSERT INTO stock_movements(rep_id,product_id,movement_type,qty,reference) VALUES (?,?,?,?,?)", (rep_id, product_id, "بيع", qty, f"فاتورة {invoice_number.strip() or order_id}"))
                    log_activity("تسجيل بيع", f"فاتورة {invoice_number.strip() or order_id} - قيمة المبيعات {total_amount:,.2f}")
                    st.success("تم اعتماد الأوردر وخصم الكمية من مخزون المندوب.")


# =========================================================
# REP INVENTORY
# =========================================================

elif page == "📦 مخزون المندوب":

    st.markdown(
        "## 📦 مخزون المندوب"
    )

    products = query(
        """
        SELECT *

        FROM products

        WHERE active = 1

        ORDER BY name
        """
    )

    data = []

    for _, product in products.iterrows():

        stock = query(
            """
            SELECT
                COALESCE(qty, 0) AS qty

            FROM rep_stock

            WHERE rep_id = ?

            AND product_id = ?
            """,
            (
                rep_id,
                int(product.id),
            ),
        )

        qty = (
            float(stock.iloc[0]["qty"])
            if len(stock)
            else 0
        )

        data.append(
            [
                product.sku,
                product.name,
                product.category,
                qty,
                product.min_stock,
            ]
        )

    inventory = pd.DataFrame(
        data,
        columns=[
            "SKU",
            "الصنف",
            "التصنيف",
            "الرصيد",
            "حد إعادة الطلب",
        ],
    )

    inventory["الحالة"] = inventory.apply(

        lambda row:

        "🔴 نفد"

        if row["الرصيد"] <= 0

        else (

            "🟠 منخفض"

            if row["الرصيد"]
            <= row["حد إعادة الطلب"]

            else "🟢 جيد"
        ),

        axis=1,
    )

    show_table(
        inventory,
        use_container_width=True,
        hide_index=True,
    )

    st.markdown("### 🔄 إضافة مخزون")
    st.caption("أي إضافة أو تعديل على مخزونك تُسجل باسمك وتظهر مباشرة في سجل الإدارة.")

    product_name = st.selectbox(
        "الصنف",
        products["name"].tolist(),
        key="rep_stock_product",
    )

    product_id = int(
        products.loc[
            products["name"]
            == product_name,
            "id",
        ].iloc[0]
    )

    quantity = st.number_input(
        "الكمية المضافة",
        min_value=0.0,
        step=1.0,
    )

    if st.button(
        "إضافة للمخزون",
        type="primary",
    ):

        execute(
            """
            INSERT INTO rep_stock
            (
                rep_id,
                product_id,
                qty
            )

            VALUES (?, ?, ?)

            ON CONFLICT
            (
                rep_id,
                product_id
            )

            DO UPDATE SET
                qty = qty + ?
            """,
            (
                rep_id,
                product_id,
                quantity,
                quantity,
            ),
        )

        execute("INSERT INTO stock_movements(rep_id,product_id,movement_type,qty,reference) VALUES (?,?,?,?,?)", (rep_id, product_id, "إضافة مخزون", quantity, "إضافة من الموظف"))
        log_activity("تحديث مخزون", f"الصنف {product_name} - الكمية {quantity:g}")
        st.success("تم تحديث مخزون المندوب.")


# =========================================================
# COLLECTIONS
# =========================================================

elif page == "💰 التحصيل":

    st.markdown(
        "## 💰 تسجيل التحصيل"
    )

    customers = query(
        """
        SELECT *

        FROM customers

        WHERE active = 1

        AND rep_id = ?

        ORDER BY name
        """,
        (rep_id,),
    )

    if len(customers):

        customer_name = st.selectbox(
            "العميل",
            customers["name"].tolist(),
        )

        customer_id = int(
            customers.loc[
                customers["name"]
                == customer_name,
                "id",
            ].iloc[0]
        )

        amount = st.number_input(
            "قيمة التحصيل",
            min_value=0.0,
            step=100.0,
        )

        reference = st.text_input(
            "مرجع / رقم إيصال"
        )

        if st.button(
            "💾 حفظ التحصيل",
            type="primary",
        ):

            execute(
                """
                INSERT INTO collections
                (
                    rep_id,
                    customer_id,
                    collection_date,
                    amount,
                    reference
                )

                VALUES (?, ?, ?, ?, ?)
                """,
                (
                    rep_id,
                    customer_id,
                    date.today().isoformat(),
                    amount,
                    reference,
                ),
            )

            st.success(
                "تم تسجيل التحصيل بنجاح."
            )

    st.markdown(
        "### آخر عمليات التحصيل"
    )

    collection_data = query(
        """
        SELECT

            c.name AS العميل,

            col.collection_date AS التاريخ,

            col.amount AS المبلغ,

            col.reference AS المرجع

        FROM collections col

        JOIN customers c
            ON c.id = col.customer_id

        WHERE col.rep_id = ?

        ORDER BY col.id DESC

        LIMIT 50
        """,
        (rep_id,),
    )

    show_table(
        collection_data,
        use_container_width=True,
        hide_index=True,
    )


# =========================================================
# REP REPORTS
# =========================================================

elif page == "📊 تقارير المندوب":

    st.markdown("## 📊 تقارير المندوب")
    st.markdown("### 📝 التقرير الأسبوعي")
    this_monday = date.today() - timedelta(days=date.today().weekday())

    with st.expander("➕ إضافة تقرير أسبوعي", expanded=True):
        template = make_weekly_report_template(current_user["name"], this_monday.isoformat(), (this_monday + timedelta(days=6)).isoformat())
        st.download_button("📄 تحميل نموذج Word", data=template or b"", file_name=f"نموذج_تقرير_{current_user['name']}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", disabled=template is None)
        week_input = st.date_input("بداية الأسبوع", this_monday, key="week_report_start")
        week_start = week_input - timedelta(days=week_input.weekday())
        week_end = week_start + timedelta(days=6)
        existing = query("SELECT id,status FROM weekly_reports WHERE rep_id=? AND week_start=? AND week_end=?", (rep_id, week_start.isoformat(), week_end.isoformat()))
        title = st.text_input("عنوان التقرير", value=f"تقرير الأسبوع {week_start.isoformat()} إلى {week_end.isoformat()}")
        summary = st.text_area("ملخص ما أنجزته خلال الأسبوع")
        file_upload = st.file_uploader("📎 إرفاق التقرير — Word (DOCX) فقط", type=["docx"], key="weekly_word_file")
        if len(existing):
            st.info(f"تم رفع تقرير هذا الأسبوع مسبقاً. الحالة: {existing.iloc[0]['status']}")
        submit = st.button("📤 إرسال التقرير الأسبوعي", type="primary", use_container_width=True, disabled=len(existing)>0)
        if submit:
            if file_upload is None:
                st.error("يجب إرفاق ملف Word بصيغة DOCX.")
            elif not summary.strip():
                st.error("اكتب ملخص التقرير أولاً.")
            else:
                path = save_upload(file_upload, "weekly_report")
                execute("INSERT INTO weekly_reports(rep_id,week_start,week_end,title,summary,file_path) VALUES (?,?,?,?,?,?)", (rep_id, week_start.isoformat(), week_end.isoformat(), title.strip(), summary.strip(), path))
                log_activity("إرسال تقرير أسبوعي", f"الأسبوع {week_start.isoformat()} إلى {week_end.isoformat()}")
                st.success("تم إرسال التقرير الأسبوعي للإدارة.")
                st.rerun()

    history = query("SELECT week_start AS بداية_الأسبوع, week_end AS نهاية_الأسبوع, title AS التقرير, status AS الحالة, manager_note AS ملاحظة_الإدارة, created_at AS تاريخ_الإرسال FROM weekly_reports WHERE rep_id=? ORDER BY week_start DESC", (rep_id,))
    show_table(history, use_container_width=True, hide_index=True)

    st.markdown("### 💰 ملخص مبيعاتي")
    my_sales = query("SELECT COUNT(*) AS عدد_الفواتير, COALESCE(SUM(total_amount),0) AS قيمة_المبيعات, COALESCE(SUM((SELECT SUM(qty) FROM order_items oi WHERE oi.order_id=o.id)),0) AS عدد_الصناديق FROM orders o WHERE o.rep_id=?", (rep_id,))
    show_table(my_sales, use_container_width=True, hide_index=True)

    st.markdown("### 📅 تقرير اليوم")
    selected_date = st.date_input(
        "التاريخ",
        date.today(),
    )

    selected_date = selected_date.isoformat()

    report = query(
        """
        SELECT

            o.order_date AS التاريخ,

            c.name AS العميل,

            o.invoice_no AS الفاتورة,

            SUM(oi.qty) AS الكمية

        FROM orders o

        JOIN customers c
            ON c.id = o.customer_id

        JOIN order_items oi
            ON oi.order_id = o.id

        WHERE o.rep_id = ?

        AND o.order_date = ?

        GROUP BY o.id

        ORDER BY o.order_date DESC
        """,
        (
            rep_id,
            selected_date,
        ),
    )

    show_table(
        report,
        use_container_width=True,
        hide_index=True,
    )

    if len(report):

        csv_data = report.to_csv(
            index=False
        ).encode("utf-8-sig")

        st.download_button(
            "⬇️ تصدير التقرير",
            csv_data,
            file_name=f"rep_report_{selected_date}.csv",
            mime="text/csv",
        )


# =========================================================
# ADMIN DASHBOARD
# =========================================================

elif page == "🏢 لوحة الإدارة":

    st.markdown("## 🏢 لوحة الإدارة")
    st.markdown(
        '<div class="app-subtitle">رؤية محدودة للمبيعات والمخزون والزيارات والتحصيل</div>',
        unsafe_allow_html=True,
    )

    today = date.today().isoformat()
    month = date.today().strftime("%Y-%m")

    # مبيعات الشهر — قيمة فعلية للفواتير
    sales = query("""
        SELECT COALESCE(SUM(oi.qty * COALESCE(oi.unit_price, p.unit_price, 0)), 0) AS amount
        FROM orders o
        JOIN order_items oi ON oi.order_id = o.id
        JOIN products p ON p.id = oi.product_id
        WHERE substr(o.order_date, 1, 7) = ?
    """, (month,))
    sales_amount = float(sales.iloc[0]["amount"]) if len(sales) else 0.0

    # المخزون الحالي — مجموع الكميات الموجودة عند جميع المندوبين
    stock = query("""
        SELECT COALESCE(SUM(qty), 0) AS qty
        FROM rep_stock
    """)
    stock_qty = float(stock.iloc[0]["qty"]) if len(stock) else 0.0

    # زيارات اليوم
    visits = query("""
        SELECT COUNT(*) AS count
        FROM visits
        WHERE visit_date = ?
    """, (today,))
    visits_count = int(visits.iloc[0]["count"]) if len(visits) else 0

    # التحصيل خلال الشهر
    collections = query("""
        SELECT COALESCE(SUM(amount), 0) AS amount
        FROM collections
        WHERE substr(collection_date, 1, 7) = ?
    """, (month,))
    collection_amount = float(collections.iloc[0]["amount"]) if len(collections) else 0.0

    # لوحة بسيطة بدون بطاقات HTML أو عرض كود
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.metric("💰 المبيعات", f"{sales_amount:,.0f} ر.س")
    with c2:
        st.metric("📦 المخزون", f"{stock_qty:,.0f} صندوق")
    with c3:
        st.metric("👥 الزيارات", f"{visits_count:,} زيارة")
    with c4:
        st.metric("💵 التحصيل", f"{collection_amount:,.0f} ر.س")

    st.divider()
    st.caption(f"ملخص الإدارة — الشهر الحالي: {month} | زيارات اليوم: {today}")


# =========================================================
# CUSTOMERS AND REPS
# =========================================================

elif page == "👥 العملاء والمندوبون":

    st.markdown(
        "## 👥 العملاء والمندوبون"
    )

    tab_reps, tab_customers = st.tabs(
        [
            "المندوبون",
            "العملاء",
        ]
    )

    # -----------------------------------------------------
    # Reps
    # -----------------------------------------------------

    with tab_reps:

        reps_data = query(
            """
            SELECT
                id,
                name,
                active

            FROM reps

            ORDER BY name
            """
        )

        show_table(
            reps_data,
            use_container_width=True,
            hide_index=True,
        )

        with st.form("add_rep"):

            name = st.text_input(
                "اسم المندوب"
            )

            submitted = st.form_submit_button(
                "إضافة مندوب"
            )

            if submitted:

                if name.strip():

                    execute(
                        """
                        INSERT INTO reps(name)
                        VALUES (?)
                        """,
                        (name,),
                    )

                    st.success(
                        "تمت إضافة المندوب."
                    )

    # -----------------------------------------------------
    # Customers
    # -----------------------------------------------------

    with tab_customers:

        customers_data = query(
            """
            SELECT

                c.id,

                c.name AS العميل,

                c.phone AS الجوال,

                c.area AS المنطقة,

                r.name AS المندوب

            FROM customers c

            LEFT JOIN reps r
                ON r.id = c.rep_id

            ORDER BY c.name
            """
        )

        show_table(
            customers_data,
            use_container_width=True,
            hide_index=True,
        )

        with st.form(
            "add_customer"
        ):

            name = st.text_input("اسم العميل")
            customer_phone = st.text_input("رقم جوال العميل")
            area = st.text_input(
                "المنطقة"
            )

            reps_list = query(
                """
                SELECT *

                FROM reps

                WHERE active = 1

                ORDER BY name
                """
            )

            if len(reps_list):

                rep_name = st.selectbox(
                    "المندوب",
                    reps_list["name"].tolist(),
                )

                submitted = st.form_submit_button(
                    "إضافة عميل"
                )

                if submitted:

                    rep_id_new = int(
                        reps_list.loc[
                            reps_list["name"]
                            == rep_name,
                            "id",
                        ].iloc[0]
                    )

                    execute(
                        """
                        INSERT INTO customers(name,area,rep_id,phone)
                        VALUES (?, ?, ?, ?)
                        """,
                        (
                            name,
                            area,
                            rep_id_new,
                            customer_phone.strip(),
                        ),
                    )

                    st.success(
                        "تمت إضافة العميل."
                    )


# =========================================================
# PRODUCTS AND INVENTORY
# =========================================================

elif page == "📦 الأصناف والمخزون":

    st.markdown(
        "## 📦 الأصناف والمخزون"
    )

    products = query(
        """
        SELECT *

        FROM products

        WHERE active = 1

        ORDER BY name
        """
    )

    show_table(
        products,
        use_container_width=True,
        hide_index=True,
    )

    st.markdown(
        "### ➕ إضافة صنف"
    )

    with st.form(
        "add_product"
    ):

        sku = st.text_input(
            "SKU"
        )

        name = st.text_input(
            "اسم الصنف"
        )

        category = st.text_input("التصنيف", value="ورق")
        unit_price = st.number_input("سعر بيع الكرتون / الوحدة", min_value=0.0, step=1.0)
        barcode = st.text_input("باركود الصنف")

        min_stock = st.number_input(
            "الحد الأدنى للمخزون",
            min_value=0,
            step=1,
        )

        dormant_days = st.number_input(
            "عدد أيام الركود",
            min_value=1,
            value=30,
            step=1,
        )

        submitted = st.form_submit_button(
            "إضافة الصنف"
        )

        if submitted:

            try:

                execute(
                    """
                    INSERT INTO products
                    (
                        sku,
                        name,
                        category,
                        min_stock,
                        dormant_days,
                        unit_price,
                        barcode
                    )

                    VALUES (?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        sku,
                        name,
                        category,
                        min_stock,
                        dormant_days,
                        unit_price,
                        barcode.strip(),
                    ),
                )

                st.success(
                    "تمت إضافة الصنف."
                )

            except Exception:

                st.error(
                    "تعذر إضافة الصنف. "
                    "تأكد من أن SKU غير مكرر."
                )

    st.markdown(
        "### 📊 مخزون جميع المندوبين"
    )

    all_stock = query(
        """
        SELECT

            r.name AS المندوب,

            p.sku,

            p.name AS الصنف,

            COALESCE(
                rs.qty,
                0
            ) AS الرصيد

        FROM reps r

        CROSS JOIN products p

        LEFT JOIN rep_stock rs

            ON rs.rep_id = r.id

            AND rs.product_id = p.id

        WHERE r.active = 1

        AND p.active = 1

        ORDER BY
            r.name,
            p.name
        """
    )

    show_table(
        all_stock,
        use_container_width=True,
        hide_index=True,
    )


# =========================================================
# TARGETS
# =========================================================

elif page == "🎯 التارجت والتحصيل":

    st.markdown(
        "## 🎯 التارجت والتحصيل"
    )

    reps_list = query(
        """
        SELECT *

        FROM reps

        WHERE active = 1

        ORDER BY name
        """
    )

    if len(reps_list):

        with st.form(
            "target_form"
        ):

            rep_name = st.selectbox(
                "المندوب",
                reps_list["name"].tolist(),
            )

            selected_rep_id = int(
                reps_list.loc[
                    reps_list["name"]
                    == rep_name,
                    "id",
                ].iloc[0]
            )

            month = st.text_input(
                "الشهر YYYY-MM",
                value=date.today().strftime(
                    "%Y-%m"
                ),
            )

            sales_target = st.number_input(
                "تارجت المبيعات",
                min_value=0.0,
                step=100.0,
            )

            collection_target = st.number_input(
                "تارجت التحصيل",
                min_value=0.0,
                step=100.0,
            )

            submitted = st.form_submit_button(
                "💾 حفظ التارجت"
            )

            if submitted:

                execute(
                    """
                    INSERT INTO targets
                    (
                        rep_id,
                        month,
                        target_sales,
                        target_collection
                    )

                    VALUES (?, ?, ?, ?)
                    """,
                    (
                        selected_rep_id,
                        month,
                        sales_target,
                        collection_target,
                    ),
                )

                st.success(
                    "تم حفظ التارجت."
                )

    targets = query(
        """
        SELECT

            r.name AS المندوب,

            t.month AS الشهر,

            t.target_sales AS تارجت_المبيعات,

            t.target_collection AS تارجت_التحصيل

        FROM targets t

        JOIN reps r
            ON r.id = t.rep_id

        ORDER BY
            t.month DESC,
            r.name
        """
    )

    show_table(
        targets,
        use_container_width=True,
        hide_index=True,
    )


# =========================================================
# DORMANT PRODUCTS
# =========================================================

elif page == "⚠️ الأصناف الراكدة":

    st.markdown(
        "## ⚠️ الأصناف الراكدة لدى العملاء"
    )

    dormant_days = st.number_input(
        "اعتبر الصنف راكداً بعد عدد أيام",
        min_value=1,
        value=30,
        step=1,
    )

    cutoff = (
        date.today()
        - timedelta(
            days=int(
                dormant_days
            )
        )
    ).isoformat()

    dormant = query(
        """
        WITH latest AS (

            SELECT

                v.customer_id,

                vs.product_id,

                MAX(
                    v.visit_date
                ) AS latest_visit

            FROM visits v

            JOIN visit_stock vs

                ON vs.visit_id = v.id

            GROUP BY
                v.customer_id,
                vs.product_id
        )

        SELECT

            c.name AS العميل,

            p.name AS الصنف,

            latest.latest_visit
                AS آخر_زيارة,

            CAST(

                julianday(?)
                - julianday(
                    latest.latest_visit
                )

                AS INTEGER

            ) AS أيام_منذ_الزيارة

        FROM latest

        JOIN customers c
            ON c.id = latest.customer_id

        JOIN products p
            ON p.id = latest.product_id

        WHERE latest.latest_visit <= ?

        ORDER BY
            أيام_منذ_الزيارة DESC
        """,
        (
            date.today().isoformat(),
            cutoff,
        ),
    )

    if len(dormant):

        st.error(
            f"تم العثور على {len(dormant)} "
            "حالة تحتاج متابعة."
        )

    else:

        st.success(
            "لا توجد أصناف راكدة حسب المدة المحددة."
        )

    show_table(
        dormant,
        use_container_width=True,
        hide_index=True,
    )


# =========================================================
# ADMIN REPORTS
# =========================================================

elif page == "📥 التقارير والتصدير":

    st.markdown(
        "## 📥 التقارير اليومية والتصدير"
    )

    report_date = st.date_input(
        "تاريخ التقرير",
        date.today(),
    ).isoformat()

    # -----------------------------------------------------
    # Visits
    # -----------------------------------------------------

    visits_report = query(
        """
        SELECT

            v.visit_date AS التاريخ,

            r.name AS المندوب,

            c.name AS العميل,

            v.notes AS الملاحظات

        FROM visits v

        JOIN reps r
            ON r.id = v.rep_id

        JOIN customers c
            ON c.id = v.customer_id

        WHERE v.visit_date = ?

        ORDER BY v.id DESC
        """,
        (report_date,),
    )

    # -----------------------------------------------------
    # Orders
    # -----------------------------------------------------

    orders_report = query(
        """
        SELECT

            o.order_date AS التاريخ,

            r.name AS المندوب,

            c.name AS العميل,

            o.invoice_no AS الفاتورة,

            SUM(
                oi.qty
            ) AS الكمية

        FROM orders o

        JOIN reps r
            ON r.id = o.rep_id

        JOIN customers c
            ON c.id = o.customer_id

        JOIN order_items oi
            ON oi.order_id = o.id

        WHERE o.order_date = ?

        GROUP BY o.id

        ORDER BY o.id DESC
        """,
        (report_date,),
    )

    # -----------------------------------------------------
    # Collections
    # -----------------------------------------------------

    collections_report = query(
        """
        SELECT

            col.collection_date AS التاريخ,

            r.name AS المندوب,

            c.name AS العميل,

            col.amount AS المبلغ,

            col.reference AS المرجع

        FROM collections col

        JOIN reps r
            ON r.id = col.rep_id

        JOIN customers c
            ON c.id = col.customer_id

        WHERE col.collection_date = ?

        ORDER BY col.id DESC
        """,
        (report_date,),
    )

    st.markdown(
        "### 👥 الزيارات"
    )

    show_table(
        visits_report,
        use_container_width=True,
        hide_index=True,
    )

    st.markdown(
        "### 🧾 الأوردرات"
    )

    show_table(
        orders_report,
        use_container_width=True,
        hide_index=True,
    )

    st.markdown(
        "### 💰 التحصيل"
    )

    show_table(
        collections_report,
        use_container_width=True,
        hide_index=True,
    )

    # -----------------------------------------------------
    # Downloads
    # -----------------------------------------------------

    st.markdown("### 📝 التقارير الأسبوعية للموظفين")
    weekly_admin = query("""
        SELECT wr.id, r.name AS الموظف, wr.week_start AS بداية_الأسبوع, wr.week_end AS نهاية_الأسبوع, wr.title AS التقرير, wr.status AS الحالة, wr.created_at AS الإرسال, wr.manager_note AS ملاحظة_الإدارة, wr.file_path
        FROM weekly_reports wr JOIN reps r ON r.id=wr.rep_id ORDER BY wr.id DESC
    """)
    if len(weekly_admin):
        show_table(weekly_admin.drop(columns=["file_path"]), use_container_width=True, hide_index=True)

    st.markdown(
        "### ⬇️ تحميل التقارير"
    )

    col1, col2, col3 = st.columns(3)

    with col1:

        st.download_button(
            "⬇️ تقرير الزيارات",
            visits_report.to_csv(
                index=False
            ).encode("utf-8-sig"),
            file_name=f"visits_{report_date}.csv",
            mime="text/csv",
        )

    with col2:

        st.download_button(
            "⬇️ تقرير الأوردرات",
            orders_report.to_csv(
                index=False
            ).encode("utf-8-sig"),
            file_name=f"orders_{report_date}.csv",
            mime="text/csv",
        )

    with col3:

        st.download_button(
            "⬇️ تقرير التحصيل",
            collections_report.to_csv(
                index=False
            ).encode("utf-8-sig"),
            file_name=f"collections_{report_date}.csv",
            mime="text/csv",
        )


# =========================================================
# ADMIN ACTIVITY
# =========================================================

elif page == "👁️ سجل نشاط الموظفين":
    st.markdown("## 👁️ سجل نشاط الموظفين")
    st.caption("يعرض للإدارة ما قام به الموظفون داخل النظام.")
    employees = query("SELECT id,name FROM users WHERE role='employee' ORDER BY name")
    filter_name = st.selectbox("الموظف", ["الكل"] + employees["name"].tolist())
    params = []
    sql = "SELECT created_at AS الوقت, user_name AS المستخدم, CASE role WHEN 'employee' THEN 'موظف' WHEN 'deputy' THEN 'نائب مدير' ELSE 'مدير' END AS الصلاحية, action AS العملية, details AS التفاصيل FROM activity_log WHERE 1=1"
    if filter_name != "الكل":
        sql += " AND user_name = ?"
        params.append(filter_name)
    sql += " ORDER BY id DESC LIMIT 500"
    activity = query(sql, tuple(params))
    show_table(activity, use_container_width=True, hide_index=True)


# =========================================================
# ADMIN EMPLOYEE FULL PROFILE
# =========================================================

elif page == "👤 ملف الموظف الكامل":
    st.markdown("## 👤 ملف الموظف الكامل")
    st.caption("هنا ترى الإدارة كل ما فعله الموظف: العملاء، الفواتير، المبيعات، الزيارات، المخزون، التقارير وسجل النشاط.")
    employees = query("SELECT id,name,rep_id FROM users WHERE role='employee' AND active=1 ORDER BY name")
    if not len(employees):
        st.info("لا يوجد موظفون فعالون.")
    else:
        employee_name = st.selectbox("اختر الموظف", employees["name"].tolist())
        emp = employees[employees["name"]==employee_name].iloc[0]
        emp_rep_id = int(emp["rep_id"])
        stats = query("""
            SELECT
                (SELECT COUNT(*) FROM customers WHERE rep_id=? AND active=1) AS العملاء,
                (SELECT COUNT(*) FROM visits WHERE rep_id=?) AS الزيارات,
                (SELECT COUNT(*) FROM orders WHERE rep_id=?) AS الفواتير,
                (SELECT COALESCE(SUM(total_amount),0) FROM orders WHERE rep_id=?) AS المبيعات,
                (SELECT COUNT(*) FROM weekly_reports WHERE rep_id=?) AS التقارير
        """, (emp_rep_id,emp_rep_id,emp_rep_id,emp_rep_id,emp_rep_id))
        cols=st.columns(5)
        for col,label in zip(cols,["العملاء","الزيارات","الفواتير","المبيعات","التقارير"]):
            value=stats.iloc[0][label]
            col.metric(label, f"{float(value):,.0f}" if label=="المبيعات" else int(value))

        tabs=st.tabs(["🧾 الفواتير والمبيعات","👥 العملاء","📝 التقارير الأسبوعية","📦 المخزون والحركة","👁️ سجل النشاط"])
        with tabs[0]:
            orders=query("""
                SELECT o.id AS رقم, o.order_date AS التاريخ, c.name AS العميل, o.invoice_no AS الفاتورة, o.invoice_barcode AS الباركود, o.total_amount AS قيمة_المبيعات, o.image_path AS صورة_الفاتورة, o.barcode_image_path AS صورة_الباركود
                FROM orders o JOIN customers c ON c.id=o.customer_id WHERE o.rep_id=? ORDER BY o.id DESC
            """, (emp_rep_id,))
            if len(orders):
                show_table(orders.drop(columns=["صورة_الفاتورة","صورة_الباركود"]),use_container_width=True,hide_index=True)
                selected_order=int(st.selectbox("عرض صورة فاتورة",orders["رقم"].tolist()))
                row=orders[orders["رقم"]==selected_order].iloc[0]
                if row["صورة_الفاتورة"] and Path(str(row["صورة_الفاتورة"])).exists():
                    st.image(str(row["صورة_الفاتورة"]),caption=f"فاتورة {row['الفاتورة'] or selected_order}")
        with tabs[1]:
            cust=query("SELECT name AS العميل, phone AS الجوال, area AS المنطقة, notes AS ملاحظات FROM customers WHERE rep_id=? ORDER BY id DESC",(emp_rep_id,))
            show_table(cust,use_container_width=True,hide_index=True)
        with tabs[2]:
            reports=query("SELECT id,week_start AS بداية_الأسبوع,week_end AS نهاية_الأسبوع,title AS التقرير,summary AS الملخص,status AS الحالة,manager_note AS ملاحظة_الإدارة,file_path FROM weekly_reports WHERE rep_id=? ORDER BY week_start DESC",(emp_rep_id,))
            if len(reports):
                show_table(reports.drop(columns=["file_path"]),use_container_width=True,hide_index=True)
                selected_report=int(st.selectbox("اختيار التقرير",reports["id"].tolist()))
                r=reports[reports["id"]==selected_report].iloc[0]
                data=file_bytes(r["file_path"])
                if data:
                    st.download_button("📄 تحميل تقرير Word",data=data,file_name=Path(str(r["file_path"])).name,mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                with st.form(f"review_report_{selected_report}"):
                    status_options=["pending","approved","rejected"]
                    current_status=str(r["الحالة"])
                    new_status=st.selectbox("حالة التقرير",status_options,index=status_options.index(current_status) if current_status in status_options else 0)
                    note=st.text_area("ملاحظة الإدارة",value="" if pd.isna(r["ملاحظة_الإدارة"]) else str(r["ملاحظة_الإدارة"]))
                    if st.form_submit_button("💾 حفظ المراجعة"):
                        execute("UPDATE weekly_reports SET status=?,manager_note=?,reviewed_by=?,reviewed_at=? WHERE id=?",(new_status,note,current_user["id"],datetime.now().isoformat(timespec="seconds"),selected_report))
                        log_activity("مراجعة تقرير أسبوعي",f"الموظف: {employee_name} - الحالة: {new_status}")
                        st.success("تم حفظ المراجعة.")
                        st.rerun()
        with tabs[3]:
            stock=query("SELECT p.name AS الصنف,p.sku AS SKU,COALESCE(rs.qty,0) AS الرصيد FROM products p LEFT JOIN rep_stock rs ON rs.product_id=p.id AND rs.rep_id=? WHERE p.active=1 ORDER BY p.name",(emp_rep_id,))
            show_table(stock,use_container_width=True,hide_index=True)
            movements=query("SELECT sm.created_at AS الوقت,p.name AS الصنف,sm.movement_type AS الحركة,sm.qty AS الكمية,sm.reference AS المرجع FROM stock_movements sm JOIN products p ON p.id=sm.product_id WHERE sm.rep_id=? ORDER BY sm.id DESC LIMIT 500",(emp_rep_id,))
            show_table(movements,use_container_width=True,hide_index=True)
        with tabs[4]:
            activity=query("SELECT created_at AS الوقت,action AS العملية,details AS التفاصيل FROM activity_log WHERE user_name=? ORDER BY id DESC LIMIT 1000",(employee_name,))
            show_table(activity,use_container_width=True,hide_index=True)


# =========================================================
# ADMIN ACCOUNT MANAGEMENT
# =========================================================

elif page == "🔐 إدارة الحسابات والصلاحيات":
    st.markdown("## 🔐 إدارة الحسابات والصلاحيات")
    st.caption("المدير ونائب المدير فقط يملكان صلاحية إدارة الحسابات.")

    tab_add, tab_manage = st.tabs(["➕ إضافة حساب", "🛠️ تعديل الحسابات"])

    with tab_add:
        account_type = st.radio(
            "ماذا تريد إضافة؟",
            ["موظف", "نائب مدير", "مدير"],
            horizontal=True,
        )
        role_map = {"موظف": "employee", "نائب مدير": "deputy", "مدير": "manager"}
        with st.form("add_account_form"):
            new_name = st.text_input("اسم الشخص")
            new_pin = st.text_input("الرمز السري", type="password", max_chars=20)
            new_pin_confirm = st.text_input("تأكيد الرمز السري", type="password", max_chars=20)
            submitted = st.form_submit_button("إنشاء الحساب", use_container_width=True)
            if submitted:
                if not new_name.strip() or not new_pin:
                    st.error("أدخل الاسم والرمز السري.")
                elif new_pin != new_pin_confirm:
                    st.error("الرمزان غير متطابقين.")
                elif len(new_pin) < 4:
                    st.error("الرمز السري يجب أن يكون 4 أرقام أو أكثر.")
                elif len(query("SELECT id FROM users WHERE name = ?", (new_name.strip(),))):
                    st.error("هذا الاسم مستخدم بالفعل.")
                else:
                    role_new = role_map[account_type]
                    new_rep_id = None
                    if role_new == "employee":
                        new_rep_id = execute("INSERT INTO reps(name, phone) VALUES (?, NULL)", (new_name.strip(),))
                    user_id = execute(
                        "INSERT INTO users(name, role, pin_hash, rep_id) VALUES (?, ?, ?, ?)",
                        (new_name.strip(), role_new, hash_pin(new_pin), new_rep_id),
                    )
                    log_activity("إنشاء حساب", f"تم إنشاء {account_type}: {new_name.strip()}")
                    st.success(f"تم إنشاء حساب {account_type} بنجاح.")
                    st.rerun()

    with tab_manage:
        users = query("SELECT id,name,role,active,rep_id FROM users ORDER BY CASE role WHEN 'manager' THEN 1 WHEN 'deputy' THEN 2 ELSE 3 END, name")
        if len(users):
            show_table(
                users.assign(الصلاحية=users["role"].map({"manager":"مدير","deputy":"نائب مدير","employee":"موظف"}), الحالة=users["active"].map({1:"فعال",0:"معطل"}))[['name','الصلاحية','الحالة']].rename(columns={'name':'الاسم'}),
                use_container_width=True, hide_index=True
            )

            selected = st.selectbox("اختر حسابًا للتعديل أو الحذف", users["name"].tolist())
            u = users[users["name"] == selected].iloc[0]

            with st.form("edit_account_form"):
                edit_name = st.text_input("الاسم", value=u["name"])
                edit_pin = st.text_input("رمز سري جديد (اتركه فارغًا لعدم تغييره)", type="password", max_chars=20)
                edit_active = st.checkbox("الحساب فعال", value=bool(u["active"]))
                save = st.form_submit_button("💾 حفظ التعديلات", use_container_width=True)
                if save:
                    if not edit_name.strip():
                        st.error("الاسم لا يمكن أن يكون فارغًا.")
                    elif edit_name.strip() != u["name"] and len(query("SELECT id FROM users WHERE name = ?", (edit_name.strip(),))):
                        st.error("الاسم الجديد مستخدم بالفعل.")
                    else:
                        if edit_pin:
                            execute("UPDATE users SET name=?, pin_hash=?, active=? WHERE id=?", (edit_name.strip(), hash_pin(edit_pin), int(edit_active), int(u["id"])))
                        else:
                            execute("UPDATE users SET name=?, active=? WHERE id=?", (edit_name.strip(), int(edit_active), int(u["id"])))
                        if u["role"] == "employee":
                            execute("UPDATE reps SET name=? WHERE id=(SELECT rep_id FROM users WHERE id=?)", (edit_name.strip(), int(u["id"])))
                        log_activity("تعديل حساب", f"تم تعديل حساب: {u['name']}")
                        st.success("تم حفظ التعديلات.")
                        st.rerun()

            # حذف الحساب نهائيًا من جدول الدخول، مع إبقاء السجلات التجارية التاريخية
            # (الفواتير/الزيارات/التقارير) حتى لا تضيع محاسبيًا. الموظف يصبح غير نشط.
            st.markdown("### 🗑️ حذف الحساب")
            st.warning(
                "حذف الحساب يمنع صاحبه من تسجيل الدخول نهائيًا. سجلات المبيعات والفواتير والزيارات والتقارير السابقة تبقى محفوظة في الشركة للتدقيق. "
                "إذا كان الحساب لموظف، يتم تعطيل ملف المندوب المرتبط به أيضًا."
            )

            is_current = int(u["id"]) == int(current_user.get("id", -1))
            active_manager_count = int(query("SELECT COUNT(*) AS c FROM users WHERE role='manager' AND active=1").iloc[0]["c"])
            can_delete = True
            if is_current:
                can_delete = False
                st.info("لا يمكن حذف الحساب الذي تستخدمه حاليًا. أنشئ الحساب البديل أولًا ثم احذف الحساب القديم.")
            elif u["role"] == "manager" and current_user.get("role") != "manager":
                can_delete = False
                st.error("حذف حساب المدير متاح للمدير الحالي فقط.")
            elif u["role"] == "manager" and active_manager_count <= 1:
                can_delete = False
                st.error("لا يمكن حذف آخر مدير نشط. أضف مديرًا بديلًا أولًا.")

            confirm_delete = st.checkbox(
                "أؤكد أنني أريد حذف هذا الحساب نهائيًا ومنع صاحبه من الدخول.",
                disabled=not can_delete,
                key=f"confirm_delete_{int(u['id'])}"
            )
            if st.button(
                "🗑️ حذف الحساب نهائيًا",
                type="secondary",
                use_container_width=True,
                disabled=(not can_delete or not confirm_delete),
                key=f"delete_account_{int(u['id'])}"
            ):
                target_name = str(u["name"])
                target_role = str(u["role"])
                target_rep_id = u["rep_id"]

                # سجل العملية قبل حذف سجل الدخول نفسه.
                log_activity(
                    "حذف حساب",
                    f"تم حذف حساب {target_role}: {target_name}. تم منع الحساب من الدخول، وحُفظت السجلات التجارية التاريخية."
                )

                execute("DELETE FROM users WHERE id=?", (int(u["id"]),))
                if target_role == "employee" and pd.notna(target_rep_id):
                    execute("UPDATE reps SET active=0 WHERE id=?", (int(target_rep_id),))

                st.success(f"تم حذف حساب {target_name} ومنعه من تسجيل الدخول.")
                st.rerun()

# =========================================================
# FOOTER
# =========================================================

st.sidebar.markdown("---")

st.sidebar.caption(
    "SalesFlow • Sales & Inventory Management"
)

st.sidebar.caption(
    datetime.now().strftime(
        "%Y-%m-%d %H:%M"
    )
)
